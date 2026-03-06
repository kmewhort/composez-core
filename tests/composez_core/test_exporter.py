"""Tests for the novel export module."""

import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock

from composez_core.exporter import (
    export_docx,
    export_epub,
    export_markdown,
)
from composez_core.narrative_map import NarrativeNode

try:
    import docx  # noqa: F401

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import ebooklib  # noqa: F401

    HAS_EBOOKLIB = True
except ImportError:
    HAS_EBOOKLIB = False


def _make_tree(tmp):
    """Build a small narrative tree with real PROSE.md files on disk.

    Returns a list of NarrativeNode (acts).
    """
    # Act 1
    act1_path = os.path.join(tmp, "act", "1 - The Beginning")
    ch1_path = os.path.join(act1_path, "chapter", "1 - Dawn")
    sc1_path = os.path.join(ch1_path, "scene", "1 - Awakening")
    sc2_path = os.path.join(ch1_path, "scene", "2 - The Road")
    ch2_path = os.path.join(act1_path, "chapter", "2 - Dusk")
    sc3_path = os.path.join(ch2_path, "scene", "1 - Sunset")

    for p in (sc1_path, sc2_path, sc3_path):
        os.makedirs(p, exist_ok=True)

    Path(os.path.join(sc1_path, "PROSE.md")).write_text(
        "The sun rose over the quiet village.\n\n"
        "Birds began to sing in the old oak tree.\n",
        encoding="utf-8",
    )
    Path(os.path.join(sc2_path, "PROSE.md")).write_text(
        "She set off down the winding road.\n",
        encoding="utf-8",
    )
    Path(os.path.join(sc3_path, "PROSE.md")).write_text(
        "The sky turned amber as evening fell.\n",
        encoding="utf-8",
    )

    # Build nodes
    sc1 = NarrativeNode("scene", 1, sc1_path, title="Awakening")
    sc2 = NarrativeNode("scene", 2, sc2_path, title="The Road")
    sc3 = NarrativeNode("scene", 1, sc3_path, title="Sunset")

    ch1_node = NarrativeNode("chapter", 1, ch1_path, title="Dawn")
    ch1_node.children = [sc1, sc2]
    ch2_node = NarrativeNode("chapter", 2, ch2_path, title="Dusk")
    ch2_node.children = [sc3]

    act1 = NarrativeNode("act", 1, act1_path, title="The Beginning")
    act1.children = [ch1_node, ch2_node]

    return [act1]


class TestExportMarkdown(unittest.TestCase):
    """Tests for markdown export."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp(prefix="export_md_test_")
        self.tree = _make_tree(self.tmp)

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_creates_file(self):
        out = os.path.join(self.tmp, "book.md")
        export_markdown(self.tree, out)
        self.assertTrue(os.path.isfile(out))

    def test_contains_act_heading(self):
        out = os.path.join(self.tmp, "book.md")
        export_markdown(self.tree, out)
        text = Path(out).read_text(encoding="utf-8")
        self.assertIn("# The Beginning", text)

    def test_contains_chapter_heading(self):
        out = os.path.join(self.tmp, "book.md")
        export_markdown(self.tree, out)
        text = Path(out).read_text(encoding="utf-8")
        self.assertIn("## Dawn", text)
        self.assertIn("## Dusk", text)

    def test_contains_prose(self):
        out = os.path.join(self.tmp, "book.md")
        export_markdown(self.tree, out)
        text = Path(out).read_text(encoding="utf-8")
        self.assertIn("sun rose over the quiet village", text)
        self.assertIn("winding road", text)
        self.assertIn("amber", text)

    def test_scene_break_between_scenes(self):
        out = os.path.join(self.tmp, "book.md")
        export_markdown(self.tree, out)
        text = Path(out).read_text(encoding="utf-8")
        # Between scene 1 and scene 2 in chapter 1 there should be a break
        self.assertIn("*", text)

    def test_empty_tree(self):
        out = os.path.join(self.tmp, "empty.md")
        export_markdown([], out)
        text = Path(out).read_text(encoding="utf-8")
        self.assertEqual(text.strip(), "")


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class TestExportDocx(unittest.TestCase):
    """Tests for DOCX export."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp(prefix="export_docx_test_")
        self.tree = _make_tree(self.tmp)

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_creates_file(self):
        out = os.path.join(self.tmp, "book.docx")
        export_docx(self.tree, out)
        self.assertTrue(os.path.isfile(out))
        # DOCX files start with PK (ZIP magic bytes)
        with open(out, "rb") as f:
            self.assertEqual(f.read(2), b"PK")

    def test_contains_act_heading(self):
        from docx import Document

        out = os.path.join(self.tmp, "book.docx")
        export_docx(self.tree, out)
        doc = Document(out)

        headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        self.assertIn("The Beginning", headings)

    def test_contains_chapter_headings(self):
        from docx import Document

        out = os.path.join(self.tmp, "book.docx")
        export_docx(self.tree, out)
        doc = Document(out)

        headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
        self.assertIn("Dawn", headings)
        self.assertIn("Dusk", headings)

    def test_contains_prose(self):
        from docx import Document

        out = os.path.join(self.tmp, "book.docx")
        export_docx(self.tree, out)
        doc = Document(out)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("sun rose", full_text)
        self.assertIn("winding road", full_text)

    def test_scene_break_present(self):
        from docx import Document

        out = os.path.join(self.tmp, "book.docx")
        export_docx(self.tree, out)
        doc = Document(out)

        texts = [p.text for p in doc.paragraphs]
        # Scene break ornament should appear between scenes in chapter 1
        self.assertTrue(any("*" in t and len(t.strip()) < 10 for t in texts))


@unittest.skipUnless(HAS_EBOOKLIB, "ebooklib not installed")
class TestExportEpub(unittest.TestCase):
    """Tests for EPUB export."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp(prefix="export_epub_test_")
        self.tree = _make_tree(self.tmp)

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_creates_file(self):
        out = os.path.join(self.tmp, "book.epub")
        export_epub(self.tree, out, title="Test Book", author="Test Author")
        self.assertTrue(os.path.isfile(out))

    def test_is_valid_zip(self):
        import zipfile

        out = os.path.join(self.tmp, "book.epub")
        export_epub(self.tree, out)
        self.assertTrue(zipfile.is_zipfile(out))

    def test_contains_act_title(self):
        import zipfile

        out = os.path.join(self.tmp, "book.epub")
        export_epub(self.tree, out)

        with zipfile.ZipFile(out) as zf:
            # ebooklib places content under EPUB/
            xhtml_names = [n for n in zf.namelist() if "act_1.xhtml" in n]
            self.assertTrue(len(xhtml_names) > 0)
            act_content = zf.read(xhtml_names[0]).decode("utf-8")
            self.assertIn("The Beginning", act_content)

    def test_contains_chapter_content(self):
        import zipfile

        out = os.path.join(self.tmp, "book.epub")
        export_epub(self.tree, out)

        with zipfile.ZipFile(out) as zf:
            xhtml_names = [n for n in zf.namelist() if "act_1_ch_1.xhtml" in n]
            self.assertTrue(len(xhtml_names) > 0)
            ch_content = zf.read(xhtml_names[0]).decode("utf-8")
            self.assertIn("Dawn", ch_content)
            self.assertIn("sun rose", ch_content)

    def test_metadata(self):
        import zipfile

        out = os.path.join(self.tmp, "book.epub")
        export_epub(self.tree, out, title="My Novel", author="Jane Doe")

        with zipfile.ZipFile(out) as zf:
            # The OPF file contains metadata
            opf_names = [n for n in zf.namelist() if n.endswith(".opf")]
            self.assertTrue(len(opf_names) > 0)
            opf = zf.read(opf_names[0]).decode("utf-8")
            self.assertIn("My Novel", opf)
            self.assertIn("Jane Doe", opf)


class TestExportCommand(unittest.TestCase):
    """Tests for the /export command in NovelCommands."""

    def setUp(self):
        self.tmp = tempfile.mkdtemp(prefix="cmd_export_test_")
        self.io = MagicMock()
        self.coder = MagicMock()
        self.coder.root = self.tmp
        self.coder.repo = None
        self.tree = _make_tree(self.tmp)

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _cmds(self):
        from composez_core.novel_commands import NovelCommands

        cmds = NovelCommands(self.io, self.coder, root=self.tmp)
        # Pre-populate the narrative map cache so it finds our tree
        cmds._narrative_map = MagicMock()
        cmds._narrative_map.get_tree.return_value = self.tree
        return cmds

    def test_export_no_args(self):
        cmds = self._cmds()
        cmds.cmd_export("")
        self.io.tool_error.assert_called()
        self.assertIn("Usage", self.io.tool_error.call_args[0][0])

    def test_export_unknown_format(self):
        cmds = self._cmds()
        cmds.cmd_export("pdf")
        self.io.tool_error.assert_called()
        self.assertIn("Unknown format", self.io.tool_error.call_args[0][0])

    def test_export_markdown_default_name(self):
        cmds = self._cmds()
        cmds.cmd_export("markdown")
        self.assertTrue(os.path.isfile(os.path.join(self.tmp, "export.md")))
        self.io.tool_output.assert_called()
        self.assertIn("export.md", self.io.tool_output.call_args[0][0])

    def test_export_markdown_custom_name(self):
        cmds = self._cmds()
        cmds.cmd_export("markdown my_book.md")
        self.assertTrue(os.path.isfile(os.path.join(self.tmp, "my_book.md")))

    @unittest.skipUnless(HAS_DOCX, "python-docx not installed")
    def test_export_docx_default_name(self):
        cmds = self._cmds()
        cmds.cmd_export("docx")
        self.assertTrue(os.path.isfile(os.path.join(self.tmp, "export.docx")))

    @unittest.skipUnless(HAS_EBOOKLIB, "ebooklib not installed")
    def test_export_epub_default_name(self):
        cmds = self._cmds()
        cmds.cmd_export("epub")
        self.assertTrue(os.path.isfile(os.path.join(self.tmp, "export.epub")))

    def test_export_empty_tree(self):
        cmds = self._cmds()
        cmds._narrative_map.get_tree.return_value = []
        cmds.cmd_export("markdown")
        self.io.tool_error.assert_called()
        self.assertIn("Nothing to export", self.io.tool_error.call_args[0][0])

    def test_export_command_registered(self):
        from composez_core.novel_commands import NovelCommands

        cmds = NovelCommands(self.io, self.coder, root=self.tmp)
        commands = cmds.get_commands()
        self.assertIn("export", commands)

    @unittest.skipUnless(HAS_EBOOKLIB, "ebooklib not installed")
    def test_export_reads_metadata(self):
        """EPUB export should pick up title/author from metadata.yml."""
        import zipfile

        # Create metadata
        core_dir = os.path.join(self.tmp, "db", "core")
        os.makedirs(core_dir, exist_ok=True)
        Path(os.path.join(core_dir, "metadata.yml")).write_text(
            "title: My Great Novel\nauthor: Jane Smith\n", encoding="utf-8"
        )

        cmds = self._cmds()
        cmds.cmd_export("epub")

        epub_path = os.path.join(self.tmp, "export.epub")
        with zipfile.ZipFile(epub_path) as zf:
            opf_names = [n for n in zf.namelist() if n.endswith(".opf")]
            opf = zf.read(opf_names[0]).decode("utf-8")
            self.assertIn("My Great Novel", opf)
            self.assertIn("Jane Smith", opf)


if __name__ == "__main__":
    unittest.main()
