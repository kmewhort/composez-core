"""
Novel export module.

Exports the narrative hierarchy into polished output formats: Markdown,
DOCX, and EPUB.  Assumes a three-level tree (top → mid → leaf) but works
with any level names configured in ``.composez``.

Formatting conventions:
    - Top-level nodes get a full-page title (page break before, centred title)
    - Mid-level nodes start on a new page with a heading
    - Leaf-level breaks use a centred ornamental separator (⁂ or "* * *")
"""

import os
from pathlib import Path


# ---------------------------------------------------------------------------
# Scene break ornament used across formats
# ---------------------------------------------------------------------------
_SCENE_BREAK_TEXT = "*\u2003*\u2003*"  # * emsp * emsp *


def _read_prose(scene_node):
    """Return the PROSE.md content for a scene node, or empty string."""
    prose_path = os.path.join(scene_node.path, "PROSE.md")
    if os.path.isfile(prose_path):
        return Path(prose_path).read_text(encoding="utf-8").strip()
    return ""


# ======================================================================
# Markdown export
# ======================================================================

def export_markdown(tree, dest_path):
    """Export the narrative tree as a single combined Markdown file.

    Parameters
    ----------
    tree : list[NarrativeNode]
        Acts returned by ``NarrativeMap.get_tree()``.
    dest_path : str
        Output file path.
    """
    lines = []

    for act in tree:
        # Act title
        act_title = act.title or f"Act {act.number}"
        lines.append(f"# {act_title}")
        lines.append("")

        for chapter in act.children:
            ch_title = chapter.title or f"Chapter {chapter.number}"
            lines.append(f"## {ch_title}")
            lines.append("")

            for i, scene in enumerate(chapter.children):
                if i > 0:
                    # Scene break between scenes within a chapter
                    lines.append("")
                    lines.append(f"<center>{_SCENE_BREAK_TEXT}</center>")
                    lines.append("")

                prose = _read_prose(scene)
                if prose:
                    lines.append(prose)
                    lines.append("")

    text = "\n".join(lines).rstrip() + "\n"
    Path(dest_path).write_text(text, encoding="utf-8")


# ======================================================================
# DOCX export
# ======================================================================

def export_docx(tree, dest_path):
    """Export the narrative tree as a styled DOCX file.

    Parameters
    ----------
    tree : list[NarrativeNode]
        Acts returned by ``NarrativeMap.get_tree()``.
    dest_path : str
        Output file path.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # -- Global defaults ---------------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5

    # -- Set page margins --------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # -- Heading styles ----------------------------------------------------
    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_font = h_style.font
        h_font.name = "Georgia"
        h_font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Heading 1: Act title (large, centred)
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(28)
    h1.font.bold = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(200)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2: Chapter title
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(12)

    # -- Build the document -----------------------------------------------
    first_act = True
    for act in tree:
        # Page break before each act
        if not first_act:
            doc.add_page_break()
        first_act = False

        # Act title page — centred with generous top spacing
        act_title = act.title or f"Act {act.number}"
        doc.add_heading(act_title, level=1)

        for ch_idx, chapter in enumerate(act.children):
            # Chapter always starts on a new page
            doc.add_page_break()

            ch_title = chapter.title or f"Chapter {chapter.number}"
            doc.add_heading(ch_title, level=2)

            for sc_idx, scene in enumerate(chapter.children):
                if sc_idx > 0:
                    _add_scene_break(doc)

                prose = _read_prose(scene)
                if prose:
                    _add_prose_paragraphs(doc, prose)

    doc.save(dest_path)


def _add_scene_break(doc):
    """Insert a centred ornamental scene break."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(_SCENE_BREAK_TEXT)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_prose_paragraphs(doc, prose_text):
    """Add prose text as paragraphs, respecting blank-line paragraph breaks."""
    from docx.shared import Pt

    paragraphs = prose_text.split("\n\n")
    for para_text in paragraphs:
        # Collapse internal single newlines into spaces
        text = " ".join(para_text.split("\n")).strip()
        if not text:
            continue
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.space_after = Pt(4)


# ======================================================================
# EPUB export
# ======================================================================

def export_epub(tree, dest_path, title="Untitled", author="Unknown"):
    """Export the narrative tree as a styled EPUB file.

    Parameters
    ----------
    tree : list[NarrativeNode]
        Acts returned by ``NarrativeMap.get_tree()``.
    dest_path : str
        Output file path.
    title : str
        Book title for EPUB metadata.
    author : str
        Author name for EPUB metadata.
    """
    import ebooklib
    from ebooklib import epub

    book = epub.EpubBook()

    # -- Metadata ----------------------------------------------------------
    book.set_identifier("novel-export")
    book.set_title(title)
    book.set_language("en")
    book.add_author(author)

    # -- Stylesheet --------------------------------------------------------
    css = _epub_stylesheet()
    style_item = epub.EpubItem(
        uid="style",
        file_name="style/default.css",
        media_type="text/css",
        content=css.encode("utf-8"),
    )
    book.add_item(style_item)

    chapters = []  # EpubHtml items
    toc = []       # table of contents entries
    spine = ["nav"]

    for act in tree:
        act_title = act.title or f"Act {act.number}"

        # Act title page
        act_page = epub.EpubHtml(
            title=act_title,
            file_name=f"act_{act.number}.xhtml",
            lang="en",
        )
        act_page.content = _epub_act_page(act_title).encode("utf-8")
        act_page.add_item(style_item)
        book.add_item(act_page)
        chapters.append(act_page)
        spine.append(act_page)

        act_toc_children = []

        for chapter in act.children:
            ch_title = chapter.title or f"Chapter {chapter.number}"

            ch_page = epub.EpubHtml(
                title=ch_title,
                file_name=f"act_{act.number}_ch_{chapter.number}.xhtml",
                lang="en",
            )
            ch_page.content = _epub_chapter_page(ch_title, chapter.children).encode("utf-8")
            ch_page.add_item(style_item)
            book.add_item(ch_page)
            chapters.append(ch_page)
            spine.append(ch_page)

            act_toc_children.append(ch_page)

        toc.append((epub.Section(act_title), act_toc_children))

    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine

    epub.write_epub(dest_path, book)


def _epub_stylesheet():
    """Return the CSS for the EPUB."""
    return """\
body {
    font-family: Georgia, "Times New Roman", serif;
    color: #1a1a1a;
    line-height: 1.6;
    margin: 1em;
}
h1.act-title {
    text-align: center;
    font-size: 2.2em;
    font-weight: bold;
    margin-top: 40%;
    margin-bottom: 0.5em;
    letter-spacing: 0.05em;
}
h2.chapter-title {
    font-size: 1.5em;
    font-weight: bold;
    margin-top: 2em;
    margin-bottom: 1em;
}
p {
    text-indent: 1.5em;
    margin: 0.25em 0;
}
p.first {
    text-indent: 0;
}
.scene-break {
    text-align: center;
    margin: 1.5em 0;
    color: #999;
    font-size: 1.2em;
    letter-spacing: 0.3em;
}
"""


def _epub_act_page(act_title):
    """Return XHTML content for an act title page."""
    safe = _esc(act_title)
    return (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<html xmlns="http://www.w3.org/1999/xhtml">\n'
        "<head><title>{t}</title>"
        '<link rel="stylesheet" href="style/default.css" type="text/css"/>'
        "</head>\n"
        '<body>\n<h1 class="act-title">{t}</h1>\n</body>\n</html>'
    ).format(t=safe)


def _epub_chapter_page(ch_title, scenes):
    """Return XHTML content for a chapter with all its scenes."""
    parts = [
        '<?xml version="1.0" encoding="utf-8"?>',
        '<html xmlns="http://www.w3.org/1999/xhtml">',
        f"<head><title>{_esc(ch_title)}</title>"
        '<link rel="stylesheet" href="style/default.css" type="text/css"/>'
        "</head>",
        "<body>",
        f'<h2 class="chapter-title">{_esc(ch_title)}</h2>',
    ]

    for i, scene in enumerate(scenes):
        if i > 0:
            parts.append(
                f'<p class="scene-break">{_esc(_SCENE_BREAK_TEXT)}</p>'
            )

        prose = _read_prose(scene)
        if prose:
            paragraphs = prose.split("\n\n")
            for j, para_text in enumerate(paragraphs):
                text = " ".join(para_text.split("\n")).strip()
                if not text:
                    continue
                cls = ' class="first"' if (i == 0 and j == 0) else ""
                parts.append(f"<p{cls}>{_esc(text)}</p>")

    parts.append("</body>")
    parts.append("</html>")
    return "\n".join(parts)


def _esc(text):
    """Escape text for safe inclusion in XHTML."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
